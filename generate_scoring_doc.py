from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# -- Styles --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# -- Title --
title = doc.add_heading('Pocket Toons — Greenlight Scoring Logic', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Comprehensive documentation of the webtoon adaptation scoring system.')

# ============================================================
doc.add_heading('1. Overview', level=1)
doc.add_paragraph(
    'The scoring system evaluates webtoon titles for adaptation potential using a '
    '3-layer decision framework that combines signals from 4 data sources: '
    'Webtoon Episodes, Webtoon Comments, YouTube Transcripts, and Reddit Posts.'
)
doc.add_paragraph(
    'Pipeline Flow:  Data Sources → Signal Extraction → Unification → '
    'Layer 1 (Hard Gates) → Layer 2 (Quality Bands) → Layer 3 (Decision Matrix) → '
    'Final Decision (GREENLIGHT / PILOT / DEFER / REWORK)'
)

# ============================================================
doc.add_heading('2. Data Sources & Signals', level=1)

t = doc.add_table(rows=7, cols=4, style='Light Grid Accent 1')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Source', 'Signal', 'Measures', 'File']
for i, h in enumerate(headers):
    t.rows[0].cells[i].text = h

data = [
    ['Webtoon Episodes', 'webtoon_cliffhanger_rate', 'Avg cliffhanger score across episodes — narrative tension', 'episode_signals.csv'],
    ['Webtoon Comments', 'webtoon_emotion_score', 'Avg emotional intensity from user comments', 'comment_signals_title.csv'],
    ['Webtoon Comments', 'webtoon_addiction_score', 'Avg addiction/binge language in comments', 'comment_signals_title.csv'],
    ['YouTube', 'youtube_hype_score', 'Audience excitement level from transcripts', 'youtube_signals.csv'],
    ['YouTube', 'youtube_confusion_score', 'Confusion/frustration in viewer reactions', 'youtube_signals.csv'],
    ['Reddit', 'reddit_risk_score', 'Sum of risk severity (pacing, fatigue, drop-off)', 'reddit_risks.csv'],
]
for r, row_data in enumerate(data, start=1):
    for c, val in enumerate(row_data):
        t.rows[r].cells[c].text = val

doc.add_paragraph()
doc.add_heading('Signal Unification', level=2)
doc.add_paragraph(
    'All signals are merged per content_id in unify_intelligence.py, producing a single '
    'content_intelligence.csv file. A data_confidence score (High/Medium/Low) tracks '
    'how many sources have data for each title.'
)

# ============================================================
doc.add_heading('3. Layer 1: Hard Gates', level=1)
doc.add_paragraph(
    'Pass/fail checks that can downgrade a decision. '
    'If a title fails any gate, its GREENLIGHT or PILOT decision is downgraded to DEFER.'
)

t2 = doc.add_table(rows=3, cols=3, style='Light Grid Accent 1')
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Gate', 'Condition to FAIL', 'Effect']):
    t2.rows[0].cells[i].text = h

t2.rows[1].cells[0].text = 'Hook Gate'
t2.rows[1].cells[1].text = 'cliffhanger_rate < 0.1 OR bottom 15% globally'
t2.rows[1].cells[2].text = 'GREENLIGHT/PILOT → DEFER'

t2.rows[2].cells[0].text = 'Clarity Gate'
t2.rows[2].cells[1].text = 'youtube_confusion_score > 0.7'
t2.rows[2].cells[2].text = 'GREENLIGHT/PILOT → DEFER'

# ============================================================
doc.add_heading('4. Layer 2: Quality Bands', level=1)
doc.add_paragraph(
    'Each signal is normalized within its genre (min-max), then classified into bands:'
)

t3 = doc.add_table(rows=4, cols=2, style='Light Grid Accent 1')
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
t3.rows[0].cells[0].text = 'Band'
t3.rows[0].cells[1].text = 'Normalized Range'
t3.rows[1].cells[0].text = 'HIGH'
t3.rows[1].cells[1].text = '≥ 0.75'
t3.rows[2].cells[0].text = 'MEDIUM'
t3.rows[2].cells[1].text = '0.25 – 0.75'
t3.rows[3].cells[0].text = 'LOW'
t3.rows[3].cells[1].text = '≤ 0.25'

doc.add_paragraph()
doc.add_heading('Composite Scores', level=2)

p = doc.add_paragraph()
p.add_run('Emotion Score (Narrative Strength):').bold = True
doc.add_paragraph('emotion = 0.4 × norm_emotion + 0.4 × norm_addiction + 0.2 × norm_cliffhanger', style='List Bullet')

p2 = doc.add_paragraph()
p2.add_run('Risk Score (Audience Frustration):').bold = True
doc.add_paragraph('risk = 0.7 × reddit_risk + 0.3 × youtube_confusion', style='List Bullet')

p3 = doc.add_paragraph()
p3.add_run('Buzz Multiplier (YouTube Hype):').bold = True
doc.add_paragraph('buzz = clamp(1.0 + 0.1 × log(1 + hype_score), min=0.9, max=1.2)', style='List Bullet')

doc.add_heading('Weight Configuration', level=2)

t4 = doc.add_table(rows=8, cols=3, style='Light Grid Accent 1')
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
t4.rows[0].cells[0].text = 'Category'
t4.rows[0].cells[1].text = 'Parameter'
t4.rows[0].cells[2].text = 'Weight'

weights = [
    ['Emotion Score', 'W_EMOTION_INTENSITY', '0.4 (40%)'],
    ['Emotion Score', 'W_ADDICTION', '0.4 (40%)'],
    ['Emotion Score', 'W_CLIFFHANGER', '0.2 (20%)'],
    ['Risk Score', 'Reddit component', '0.7 (70%)'],
    ['Risk Score', 'YouTube confusion component', '0.3 (30%)'],
    ['Buzz Multiplier', 'BUZZ_CAP_MIN', '0.9'],
    ['Buzz Multiplier', 'BUZZ_CAP_MAX', '1.2'],
]
for r, row_data in enumerate(weights, start=1):
    for c, val in enumerate(row_data):
        t4.rows[r].cells[c].text = val

# ============================================================
doc.add_heading('5. Layer 3: Decision Matrix', level=1)
doc.add_paragraph(
    'Combines overall quality band (derived from Emotion × Addiction bands) with risk band:'
)

t5 = doc.add_table(rows=4, cols=4, style='Light Grid Accent 1')
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
t5.rows[0].cells[0].text = 'Quality \\ Risk'
t5.rows[0].cells[1].text = 'LOW Risk'
t5.rows[0].cells[2].text = 'MEDIUM Risk'
t5.rows[0].cells[3].text = 'HIGH Risk'

t5.rows[1].cells[0].text = 'HIGH Quality'
t5.rows[1].cells[1].text = '✅ GREENLIGHT'
t5.rows[1].cells[2].text = '🟡 PILOT'
t5.rows[1].cells[3].text = '❌ REWORK'

t5.rows[2].cells[0].text = 'MEDIUM Quality'
t5.rows[2].cells[1].text = '🟡 PILOT'
t5.rows[2].cells[2].text = '⏸️ DEFER'
t5.rows[2].cells[3].text = '❌ REWORK'

t5.rows[3].cells[0].text = 'LOW Quality'
t5.rows[3].cells[1].text = '❌ REWORK'
t5.rows[3].cells[2].text = '❌ REWORK'
t5.rows[3].cells[3].text = '❌ REWORK'

doc.add_paragraph()
doc.add_paragraph(
    'Quality determination: HIGH = both Emotion and Addiction bands are HIGH. '
    'LOW = either Emotion or Addiction band is LOW. MEDIUM = everything else.'
)

# ============================================================
doc.add_heading('6. Decision Definitions', level=1)

decisions = [
    ('GREENLIGHT', 'Title is strongly recommended for adaptation. High narrative quality with low audience risk.'),
    ('PILOT', 'Title shows promise but needs further validation. Consider a limited pilot/test run.'),
    ('DEFER', 'Title does not meet thresholds currently. Reassess after gathering more data or improving signals.'),
    ('REWORK', 'Title has significant quality or risk issues. Not recommended for adaptation in current form.'),
]
for label, desc in decisions:
    p = doc.add_paragraph()
    p.add_run(f'{label}: ').bold = True
    p.add_run(desc)

# ============================================================
doc.add_heading('7. Current Results', level=1)

t6 = doc.add_table(rows=6, cols=4, style='Light Grid Accent 1')
t6.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Title', 'Genre', 'Decision', 'Key Reasoning']):
    t6.rows[0].cells[i].text = h

results = [
    ['Solo Leveling', 'Action', '⏸️ DEFER', 'MEDIUM quality, MEDIUM risk'],
    ['Lookism', 'Drama', '⏸️ DEFER', 'MEDIUM quality, MEDIUM risk'],
    ['Omniscient Reader', 'Fantasy', '❌ REWORK', 'HIGH emotion but LOW addiction'],
    ['Tower of God', 'Fantasy', '❌ REWORK', 'LOW emotion, HIGH risk'],
    ['True Beauty', 'Romance', '⏸️ DEFER', 'MEDIUM quality, failed Hook Gate'],
]
for r, row_data in enumerate(results, start=1):
    for c, val in enumerate(row_data):
        t6.rows[r].cells[c].text = val

doc.add_paragraph()
doc.add_paragraph(
    'Note: No titles currently qualify for GREENLIGHT. With only 5 titles, '
    'genre normalization creates heavy band compression. Expanding the dataset '
    'should improve differentiation.'
)

# ============================================================
doc.add_heading('8. Source Files', level=1)

t7 = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
t7.alignment = WD_TABLE_ALIGNMENT.CENTER
t7.rows[0].cells[0].text = 'Component'
t7.rows[0].cells[1].text = 'File'

files = [
    ['Greenlight Scorer', 'scoring/greenlight_score.py'],
    ['Signal Unifier', 'nlp/unify_intelligence.py'],
    ['Comment NLP', 'nlp/comment_signals.py'],
    ['YouTube NLP', 'nlp/youtube_signals.py'],
    ['Output', 'data/processed/greenlight_ranking.csv'],
]
for r, row_data in enumerate(files, start=1):
    for c, val in enumerate(row_data):
        t7.rows[r].cells[c].text = val

# Save
output_path = r'c:\Users\Nitin\Documents\pocket_toons\Pocket_Toons_Scoring_Logic.docx'
doc.save(output_path)
print(f"Saved to: {output_path}")
