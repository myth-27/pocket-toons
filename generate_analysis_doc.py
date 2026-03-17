"""Generate comprehensive Pocket Toons project analysis Word document."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ===================== TITLE PAGE =====================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Pocket Toons', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading('Webtoon Intelligence & Adaptation Pipeline', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
meta = doc.add_paragraph('Complete Product Analysis — Start to Finish')
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.runs[0].font.size = Pt(14)
meta.runs[0].font.color.rgb = RGBColor(100, 100, 100)
doc.add_paragraph()
date_line = doc.add_paragraph('February 28, 2026')
date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# ===================== 1. EXECUTIVE SUMMARY =====================
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'Pocket Toons is an end-to-end data intelligence platform that evaluates webtoon titles for '
    'adaptation potential (anime, live-action, etc.). The system scrapes multi-source data from '
    'Webtoons, YouTube, and Reddit, extracts NLP signals (emotion, addiction, cliffhanger tension, '
    'hype, risk), unifies them into a content intelligence table, and runs a 3-layer scoring '
    'framework to produce actionable GREENLIGHT / PILOT / DEFER / REWORK recommendations.'
)
doc.add_paragraph(
    'Additionally, an OCR pipeline has been built to extract dialogue scripts directly from webtoon '
    'comic panels using EasyOCR, creating a structured ML-ready corpus for future machine learning '
    'applications such as script generation, dialogue analysis, and content evaluation.'
)

# Key stats
doc.add_heading('Key Metrics', level=2)
t = doc.add_table(rows=8, cols=2, style='Light Grid Accent 1')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
stats = [
    ['Metric', 'Value'],
    ['Total Titles Indexed', '93'],
    ['Genres Covered', '16 (action, fantasy, romance, drama, thriller, horror, sci-fi, comedy, mystery, superhero, sports, slice of life, etc.)'],
    ['Data Sources', '4 (Webtoons, YouTube, Reddit, OCR)'],
    ['Titles Fully Scored (Greenlight)', '5'],
    ['OCR Episodes Processed', '36 (across 4 titles)'],
    ['OCR Words Extracted', '36,904'],
    ['Pipeline Components', '25+ Python scripts'],
]
for r, row_data in enumerate(stats):
    for c, val in enumerate(row_data):
        t.rows[r].cells[c].text = val

doc.add_page_break()

# ===================== 2. SYSTEM ARCHITECTURE =====================
doc.add_heading('2. System Architecture', level=1)
doc.add_paragraph(
    'The platform is organized into 6 major layers, each feeding into the next:'
)
layers = [
    ('Layer 1: Data Collection (Scrapers)', 
     'Webtoon episode/comment scrapers, YouTube transcript scraper, Reddit post scraper. '
     'Collects raw data from public APIs and web pages.'),
    ('Layer 2: Data Preparation', 
     'Cleans raw scraped data, normalizes titles, resolves content IDs, and produces '
     'structured CSVs for downstream NLP processing.'),
    ('Layer 3: NLP Signal Extraction', 
     'Extracts quantitative signals from text: emotional intensity, addiction language, '
     'cliffhanger tension, YouTube hype/confusion, Reddit risk severity.'),
    ('Layer 4: Unification', 
     'Merges all signals per content_id into a single unified intelligence table with '
     'data confidence scoring (High/Medium/Low).'),
    ('Layer 5: Greenlight Scoring', 
     '3-layer decision framework: Hard Gates → Quality Bands → Decision Matrix. '
     'Produces GREENLIGHT / PILOT / DEFER / REWORK per title.'),
    ('Layer 6: OCR Pipeline', 
     'Scrapes webtoon panel images, downloads them, runs OCR (EasyOCR), and produces '
     'structured dialogue scripts for ML training.'),
]
for name, desc in layers:
    p = doc.add_paragraph()
    p.add_run(name + ': ').bold = True
    p.add_run(desc)

doc.add_page_break()

# ===================== 3. DATA SOURCES =====================
doc.add_heading('3. Data Collection Layer', level=1)

doc.add_heading('3.1 Scrapers Built', level=2)
t2 = doc.add_table(rows=9, cols=3, style='Light Grid Accent 1')
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
scraper_data = [
    ['Scraper', 'File', 'Purpose'],
    ['Webtoon Series', 'webtoon_series_scraper.py', 'Scrapes episode list metadata for a series'],
    ['Webtoon Episodes', 'webtoon_episode_scraper.py', 'Scrapes panel images and text from individual episodes'],
    ['Webtoon Episodes (Batch)', 'webtoon_episodes_scraper.py', 'Batch scraper for multiple episodes'],
    ['Webtoon Comments', 'webtoon_comments_scraper.py', 'Extracts user comments for sentiment analysis'],
    ['Webtoon Panel Downloader', 'webtoon_panel_downloader.py', 'Downloads high-res comic panels (data-url fix)'],
    ['YouTube Transcripts', 'youtube_webtoon_transcripts.py', 'Scrapes video transcripts for hype/confusion signals'],
    ['Reddit Posts', 'reddit_scraper.py', 'Collects Reddit discussion data for risk analysis'],
]
for r, row_data in enumerate(scraper_data):
    for c, val in enumerate(row_data):
        t2.rows[r].cells[c].text = val

doc.add_heading('3.2 Raw Data Collected', level=2)
t3 = doc.add_table(rows=5, cols=3, style='Light Grid Accent 1')
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
raw_data = [
    ['File', 'Records', 'Description'],
    ['webtoon_episodes.jsonl', '9 episodes', 'Tower of God episode metadata (ep 227-235)'],
    ['webtoon_youtube_transcripts.jsonl', '120 videos', 'YouTube transcripts across 16 titles'],
    ['reddit_posts.jsonl', '35 posts', 'Reddit posts from r/manga, r/anime, r/webtoons'],
    ['webtoon_comments_test.jsonl', '~50 comments', 'Webtoon user comments sample'],
]
for r, row_data in enumerate(raw_data):
    for c, val in enumerate(row_data):
        t3.rows[r].cells[c].text = val

doc.add_page_break()

# ===================== 4. NLP SIGNAL EXTRACTION =====================
doc.add_heading('4. NLP Signal Extraction', level=1)

doc.add_heading('4.1 Webtoon Comment Signals', level=2)
doc.add_paragraph('File: nlp/comment_signals.py')
doc.add_paragraph(
    'Analyzes user comments to extract emotional intensity and addiction language. '
    'Uses keyword dictionaries and NLP patterns to score comments on a 0-1 scale.'
)
t4 = doc.add_table(rows=3, cols=3, style='Light Grid Accent 1')
t4.rows[0].cells[0].text = 'Signal'
t4.rows[0].cells[1].text = 'Range'
t4.rows[0].cells[2].text = 'What It Measures'
t4.rows[1].cells[0].text = 'webtoon_emotion_score'
t4.rows[1].cells[1].text = '0.0 – 1.0'
t4.rows[1].cells[2].text = 'Emotional intensity from comments (excitement, shock, sadness)'
t4.rows[2].cells[0].text = 'webtoon_addiction_score'
t4.rows[2].cells[1].text = '0.0 – 1.0'
t4.rows[2].cells[2].text = 'Binge/addiction language (can\'t stop, need more, hooked)'

doc.add_heading('4.2 YouTube Signals', level=2)
doc.add_paragraph('File: nlp/youtube_signals.py')
doc.add_paragraph(
    'Processes YouTube video transcripts to extract audience excitement and confusion levels. '
    'Measures community sentiment toward adaptation quality.'
)
t5 = doc.add_table(rows=3, cols=3, style='Light Grid Accent 1')
t5.rows[0].cells[0].text = 'Signal'
t5.rows[0].cells[1].text = 'Range'
t5.rows[0].cells[2].text = 'What It Measures'
t5.rows[1].cells[0].text = 'youtube_hype_score'
t5.rows[1].cells[1].text = '0.0 – 1.0'
t5.rows[1].cells[2].text = 'Audience excitement, anticipation, positive buzz'
t5.rows[2].cells[0].text = 'youtube_confusion_score'
t5.rows[2].cells[1].text = '0.0 – 1.0'
t5.rows[2].cells[2].text = 'Viewer confusion, frustration, negative signals'

doc.add_heading('4.3 Reddit Risk Signals', level=2)
doc.add_paragraph('File: nlp/reddit_risk_signals.py')
doc.add_paragraph(
    'Scans Reddit posts for risk indicators: pacing complaints, reader fatigue, '
    'drop-off mentions, confusion about plot. Aggregated as severity sum per title.'
)

doc.add_heading('4.4 Episode Cliffhanger Analysis', level=2)
doc.add_paragraph('Files: prepare_episodes.py, generate_signals.py')
doc.add_paragraph(
    'Analyzes episode metadata to score narrative tension and cliffhanger effectiveness. '
    'Higher scores indicate stronger reader retention hooks.'
)

doc.add_page_break()

# ===================== 5. UNIFIED INTELLIGENCE =====================
doc.add_heading('5. Unified Content Intelligence', level=1)
doc.add_paragraph('File: nlp/unify_intelligence.py → data/unified/content_intelligence.csv')
doc.add_paragraph(
    'Merges all extracted signals into a single row per title. Currently contains 93 titles '
    'across 16 genres with full signal coverage.'
)

doc.add_heading('Top 10 Titles by Addiction Score', level=2)
top_titles = [
    ['Title', 'Genre', 'Emotion', 'Addiction', 'Cliffhanger', 'Risk'],
    ['Bastard', 'Thriller', '0.82', '0.92', '0.91', '9.2'],
    ['Everything Is Fine', 'Thriller', '0.83', '0.91', '0.91', '4.5'],
    ['Deaths Game', 'Thriller', '0.84', '0.92', '0.89', '12.5'],
    ['Gremoryland', 'Thriller', '0.83', '0.92', '0.89', '2.4'],
    ['Homesick', 'Thriller', '0.83', '0.91', '0.92', '9.7'],
    ['Dear Diary', 'Mystery', '0.84', '0.91', '0.94', '1.7'],
    ['Dreaming Freedom', 'Mystery', '0.81', '0.88', '0.94', '9.3'],
    ['The Boxer', 'Action', '0.81', '0.88', '0.82', '10.3'],
    ['Adv. Player Tutorial Tower', 'Action', '0.85', '0.87', '0.79', '2.8'],
    ['Life Outside the Circle', 'Slice of Life', '0.83', '0.88', '0.84', '3.0'],
]
t6 = doc.add_table(rows=11, cols=6, style='Light Grid Accent 1')
t6.alignment = WD_TABLE_ALIGNMENT.CENTER
for r, row_data in enumerate(top_titles):
    for c, val in enumerate(row_data):
        t6.rows[r].cells[c].text = val

doc.add_page_break()

# ===================== 6. GREENLIGHT SCORING =====================
doc.add_heading('6. Greenlight Scoring Framework', level=1)

doc.add_heading('6.1 Layer 1: Hard Gates', level=2)
doc.add_paragraph(
    'Pass/fail checks that can downgrade decisions. If a title fails any gate, '
    'its GREENLIGHT or PILOT verdict is downgraded to DEFER.'
)
t7 = doc.add_table(rows=3, cols=3, style='Light Grid Accent 1')
t7.rows[0].cells[0].text = 'Gate'
t7.rows[0].cells[1].text = 'Fail Condition'
t7.rows[0].cells[2].text = 'Effect'
t7.rows[1].cells[0].text = 'Hook Gate'
t7.rows[1].cells[1].text = 'cliffhanger_rate < 0.1 OR bottom 15%'
t7.rows[1].cells[2].text = 'Downgrade to DEFER'
t7.rows[2].cells[0].text = 'Clarity Gate'
t7.rows[2].cells[1].text = 'youtube_confusion_score > 0.7'
t7.rows[2].cells[2].text = 'Downgrade to DEFER'

doc.add_heading('6.2 Layer 2: Quality Bands', level=2)
doc.add_paragraph('Signals are normalized within each genre (min-max), then classified:')
doc.add_paragraph('• HIGH: ≥ 0.75  |  MEDIUM: 0.25–0.75  |  LOW: ≤ 0.25')

doc.add_heading('Scoring Weights', level=3)
t8 = doc.add_table(rows=4, cols=2, style='Light Grid Accent 1')
t8.rows[0].cells[0].text = 'Component'
t8.rows[0].cells[1].text = 'Weight'
t8.rows[1].cells[0].text = 'Emotion Intensity'
t8.rows[1].cells[1].text = '40%'
t8.rows[2].cells[0].text = 'Addiction Language'
t8.rows[2].cells[1].text = '40%'
t8.rows[3].cells[0].text = 'Cliffhanger Rate'
t8.rows[3].cells[1].text = '20%'

doc.add_heading('6.3 Layer 3: Decision Matrix', level=2)
t9 = doc.add_table(rows=4, cols=4, style='Light Grid Accent 1')
t9.rows[0].cells[0].text = 'Quality \\ Risk'
t9.rows[0].cells[1].text = 'LOW Risk'
t9.rows[0].cells[2].text = 'MEDIUM Risk'
t9.rows[0].cells[3].text = 'HIGH Risk'
t9.rows[1].cells[0].text = 'HIGH Quality'
t9.rows[1].cells[1].text = 'GREENLIGHT'
t9.rows[1].cells[2].text = 'PILOT'
t9.rows[1].cells[3].text = 'REWORK'
t9.rows[2].cells[0].text = 'MEDIUM Quality'
t9.rows[2].cells[1].text = 'PILOT'
t9.rows[2].cells[2].text = 'DEFER'
t9.rows[2].cells[3].text = 'REWORK'
t9.rows[3].cells[0].text = 'LOW Quality'
t9.rows[3].cells[1].text = 'REWORK'
t9.rows[3].cells[2].text = 'REWORK'
t9.rows[3].cells[3].text = 'REWORK'

doc.add_heading('6.4 Current Scored Results', level=2)
t10 = doc.add_table(rows=6, cols=4, style='Light Grid Accent 1')
t10.alignment = WD_TABLE_ALIGNMENT.CENTER
results_data = [
    ['Title', 'Genre', 'Decision', 'Reasoning'],
    ['Solo Leveling', 'Action', 'DEFER', 'MEDIUM quality, MEDIUM risk'],
    ['Lookism', 'Drama', 'DEFER', 'MEDIUM quality, MEDIUM risk'],
    ['Omniscient Reader', 'Fantasy', 'REWORK', 'HIGH emotion / LOW addiction'],
    ['Tower of God', 'Fantasy', 'REWORK', 'LOW emotion, HIGH risk'],
    ['True Beauty', 'Romance', 'DEFER', 'MEDIUM quality, Hook Gate failed'],
]
for r, row_data in enumerate(results_data):
    for c, val in enumerate(row_data):
        t10.rows[r].cells[c].text = val

doc.add_page_break()

# ===================== 7. OCR PIPELINE =====================
doc.add_heading('7. OCR Pipeline (Webtoon Script Extraction)', level=1)

doc.add_heading('7.1 Pipeline Architecture', level=2)
steps = [
    ('Step 1: Scrape', 'Fetch episode HTML → extract panel image URLs (uses data-url attribute fix)'),
    ('Step 2: Download', 'Download all panel images to local storage (avg 100-800 panels per episode)'),
    ('Step 3: OCR', 'Run EasyOCR on each panel image to extract visible text (dialogue, narration, SFX)'),
    ('Step 4: Consolidate', 'Merge all panel text into a single episode script + detailed per-panel JSON'),
]
for step, desc in steps:
    p = doc.add_paragraph()
    p.add_run(step + ': ').bold = True
    p.add_run(desc)

doc.add_heading('7.2 OCR Strategy (Multi-Fallback)', level=2)
t11 = doc.add_table(rows=4, cols=3, style='Light Grid Accent 1')
t11.rows[0].cells[0].text = 'Strategy'
t11.rows[0].cells[1].text = 'Method'
t11.rows[0].cells[2].text = 'Status'
t11.rows[1].cells[0].text = 'A'
t11.rows[1].cells[1].text = 'Google Cloud Vision (service account)'
t11.rows[1].cells[2].text = 'Not configured — no credentials'
t11.rows[2].cells[0].text = 'B'
t11.rows[2].cells[1].text = 'Google Vision REST API (API key)'
t11.rows[2].cells[2].text = '403 — billing not enabled'
t11.rows[3].cells[0].text = 'C'
t11.rows[3].cells[1].text = 'EasyOCR (local, PyTorch CPU)'
t11.rows[3].cells[2].text = 'ACTIVE — working'

doc.add_heading('7.3 Batch Pipeline', level=2)
doc.add_paragraph('File: batch_ocr_pipeline.py')
doc.add_paragraph(
    'Orchestrates OCR across multiple titles and episodes. Resume-safe (skips already-processed '
    'episodes). Outputs structured ML-ready data.'
)

doc.add_heading('7.4 Title Registry', level=2)
t12 = doc.add_table(rows=11, cols=4, style='Light Grid Accent 1')
t12.alignment = WD_TABLE_ALIGNMENT.CENTER
registry_data = [
    ['UID', 'Title', 'Genre', 'Webtoon ID'],
    ['WT001', 'Tower of God', 'Fantasy', '95'],
    ['WT002', 'Solo Leveling', 'Action', '3162'],
    ['WT003', 'Omniscient Reader', 'Fantasy', '2154'],
    ['WT004', 'Lookism', 'Drama', '1049'],
    ['WT005', 'True Beauty', 'Romance', '1436'],
    ['WT006', 'God of High School', 'Action', '66'],
    ['WT007', 'Noblesse', 'Action', '87'],
    ['WT008', 'Sweet Home', 'Horror', '1285'],
    ['WT009', 'unOrdinary', 'Action', '679'],
    ['WT010', 'Lore Olympus', 'Romance', '1320'],
]
for r, row_data in enumerate(registry_data):
    for c, val in enumerate(row_data):
        t12.rows[r].cells[c].text = val

doc.add_heading('7.5 Current OCR Progress', level=2)
t13 = doc.add_table(rows=5, cols=3, style='Light Grid Accent 1')
t13.alignment = WD_TABLE_ALIGNMENT.CENTER
ocr_progress = [
    ['Title', 'Episodes Done', 'Status'],
    ['Tower of God', '10 / 10', 'COMPLETE'],
    ['Solo Leveling', '10 / 10', 'COMPLETE'],
    ['Omniscient Reader', '10 / 10', 'COMPLETE'],
    ['Lookism', '6 / 10', 'IN PROGRESS'],
]
for r, row_data in enumerate(ocr_progress):
    for c, val in enumerate(row_data):
        t13.rows[r].cells[c].text = val

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Total: ').bold = True
p.add_run('36 episodes processed, 36,904 words extracted')

doc.add_heading('7.6 ML Corpus Output Format', level=2)
doc.add_paragraph('Output: data/ml_dataset/ocr_corpus.jsonl')
doc.add_paragraph('Each line is a JSON object with the following structure:')
doc.add_paragraph(
    '{"uid": "WT001", "content_id": "tower_of_god", "title": "Tower of God", '
    '"genre": "fantasy", "episode": 1, "word_count": 523, "script": "..."}'
)

doc.add_page_break()

# ===================== 8. STREAMLIT DASHBOARD =====================
doc.add_heading('8. Streamlit Dashboard', level=1)
doc.add_paragraph('File: streamlit_app.py (28 KB)')
doc.add_paragraph(
    'Interactive web dashboard built with Streamlit for exploring results. Features include:'
)
features = [
    'Title-level scoring overview with decision labels',
    'Genre breakdown and filtering',
    'Signal visualization (emotion, addiction, cliffhanger, risk)',
    'Script evaluation with feedback system',
    'Video preview generation for content pitches',
    'Dynamic sidebar with global metrics (total titles, evaluations)',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# ===================== 9. FILE STRUCTURE =====================
doc.add_heading('9. Project File Structure', level=1)

structure = [
    ('config/', 'Configuration files'),
    ('  settings.yaml', 'API keys, scraping settings, paths'),
    ('  title_registry.json', 'Webtoon title registry with unique IDs'),
    ('scrapers/', 'Data collection layer (8 scrapers)'),
    ('nlp/', 'NLP signal extraction (5 modules)'),
    ('scoring/', 'Greenlight scoring framework'),
    ('models/', 'ML models and OCR handlers'),
    ('data/raw/', 'Raw scraped data (JSONL files, panel images)'),
    ('data/processed/', 'Processed signals and rankings (CSV)'),
    ('data/unified/', 'Unified content intelligence table'),
    ('data/ml_dataset/', 'ML-ready corpus output'),
    ('ocr_pipeline.py', 'Single-episode OCR pipeline'),
    ('batch_ocr_pipeline.py', 'Batch OCR orchestrator (all titles)'),
    ('streamlit_app.py', 'Interactive dashboard'),
]
t14 = doc.add_table(rows=len(structure)+1, cols=2, style='Light Grid Accent 1')
t14.rows[0].cells[0].text = 'Path'
t14.rows[0].cells[1].text = 'Purpose'
for r, (path, purpose) in enumerate(structure, start=1):
    t14.rows[r].cells[0].text = path
    t14.rows[r].cells[1].text = purpose

doc.add_page_break()

# ===================== 10. TECHNICAL CHALLENGES =====================
doc.add_heading('10. Technical Challenges Solved', level=1)

challenges = [
    ('Webtoon Image Placeholders', 
     'Webtoons uses data-url attributes (not src) for actual panel images. The src attribute '
     'contains a transparent placeholder GIF. Fixed by prioritizing data-url in the scraper.'),
    ('Google Cloud Vision API', 
     'Multiple API keys tested — all resulted in 403 errors due to: (1) billing not enabled, '
     '(2) Vision API not enabled on the project, or (3) key explicitly blocked. Resolved by '
     'implementing EasyOCR as a local CPU-based fallback.'),
    ('Tesseract OCR Installation', 
     'Tesseract binary could not be installed on Windows via winget (403 download error). '
     'Bypassed by using EasyOCR (pure Python, no system binary needed).'),
    ('Unicode Console Encoding', 
     'Windows console threw encoding errors for Korean/Japanese characters in scraped content. '
     'Fixed by reconfiguring stdout/stderr to UTF-8.'),
    ('OCR Processing Speed', 
     'EasyOCR on CPU takes ~15-20 min per episode (100-800 panels each). Made the batch pipeline '
     'resume-safe so it can be interrupted and continued.'),
]
for title, desc in challenges:
    p = doc.add_paragraph()
    p.add_run(title + ': ').bold = True
    p.add_run(desc)

doc.add_page_break()

# ===================== 11. NEXT STEPS =====================
doc.add_heading('11. Next Steps & Recommendations', level=1)

doc.add_heading('Immediate', level=2)
immediate = [
    'Complete the batch OCR pipeline (remaining 64 episodes across 6 titles)',
    'Re-run Greenlight scoring with all 93 titles (currently only 5 are scored)',
    'Enable Google Cloud Vision billing for higher-quality OCR output',
]
for item in immediate:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Short-Term', level=2)
short_term = [
    'Build ML models using OCR corpus for dialogue quality scoring',
    'Integrate OCR-derived signals into the Greenlight scoring framework',
    'Add script structure analysis (dialogue ratio, pacing, character density)',
    'Expand Reddit scraping with real API data (PRAW) for richer risk signals',
]
for item in short_term:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Long-Term', level=2)
long_term = [
    'Train transformer-based script evaluator on OCR corpus',
    'Predictive model for adaptation success based on all signals',
    'Automated content recommendation engine for producers',
    'Real-time monitoring dashboard for new episode releases',
]
for item in long_term:
    doc.add_paragraph(item, style='List Bullet')

# ===================== 12. DEPENDENCIES =====================
doc.add_heading('12. Technology Stack', level=1)
t15 = doc.add_table(rows=12, cols=2, style='Light Grid Accent 1')
tech_data = [
    ['Technology', 'Purpose'],
    ['Python 3.14', 'Core language'],
    ['BeautifulSoup4', 'HTML parsing for scrapers'],
    ['Pandas / NumPy', 'Data processing and signal calculation'],
    ['NLTK / spaCy', 'NLP text analysis'],
    ['scikit-learn', 'ML utilities and normalization'],
    ['EasyOCR (PyTorch)', 'Comic panel text extraction'],
    ['python-docx', 'Report generation'],
    ['Streamlit', 'Interactive dashboard'],
    ['Requests', 'HTTP client for scraping'],
    ['PyYAML', 'Configuration management'],
]
for r, row_data in enumerate(tech_data):
    for c, val in enumerate(row_data):
        t15.rows[r].cells[c].text = val

# Save
output_path = r'c:\Users\Nitin\Documents\pocket_toons\Pocket_Toons_Complete_Analysis.docx'
doc.save(output_path)
print(f"Saved to: {output_path}")
